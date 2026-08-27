"""Presenter briefing Word doc. Run: python presentation/build_research_doc.py"""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x1C, 0x1A, 0x16)
MOSS = RGBColor(0x3A, 0x4A, 0x38)
MUTED = RGBColor(0x6B, 0x64, 0x5A)
OUT = Path(__file__).resolve().parent / "For-the-Presenter.docx"
OLD = Path(__file__).resolve().parent / "Research-First.docx"


def set_run(run, *, size=11, bold=False, color=INK, italic=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def add_para(doc, text, *, size=11, bold=False, italic=False, color=INK, name="Calibri", space_before=0, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color, name=name)
    return p


def add_heading(doc, text, level=1):
    return add_para(
        doc,
        text,
        size=18 if level == 1 else 13,
        bold=True,
        color=MOSS if level == 1 else INK,
        name="Georgia",
        space_before=16 if level == 1 else 12,
        space_after=8,
    )


def add_bullet(doc, text, *, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run, size=size)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=INK, size=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        shade_cell(cell, "3A4A38")
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.rows[r].cells[c]
            set_cell_text(cell, value, size=10)
            if r % 2 == 0:
                shade_cell(cell, "F6F4F0")
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    add_para(doc, "FOR YOU  ·  READ BEFORE WE SIT DOWN", size=10, bold=True, color=MOSS, space_after=4)
    add_para(doc, "What you are presenting", size=26, bold=True, name="Georgia", space_after=4)
    add_para(
        doc,
        "You will give a short talk about a marketing website. You do not need to have written the code. "
        "Read this once, slowly. Click the sites if you can. Then we will mentor: the slides, the demo, and the speaking script. "
        "This briefing is so we are not starting from zero.",
        size=12,
        italic=True,
        color=MUTED,
        name="Georgia",
        space_after=12,
    )

    add_heading(doc, "The talk in one page")
    add_para(
        doc,
        "You are presenting a marketing website — not a SaaS product.",
        size=14,
        italic=True,
        name="Georgia",
        space_after=10,
    )
    add_para(
        doc,
        "SaaS is an app people log into: Slack, Notion, a dashboard with settings and billing. The software is the product. "
        "This talk is the opposite. Nobody logs in. There is no account. It is a public website for a real business — "
        "a Hudson Valley farm that sells heirloom seed. The pages are the catalog. Marketers change packets and headings "
        "in a CMS. Engineering does not ship a new website build for those changes.",
    )
    add_para(
        doc,
        "The site is called Thorn & Furrow. Catalog Number 14, Spring 2026. Visitors browse packets, read the journal, "
        "and request a packet list by mail. That request is a marketing lead — not a SaaS signup.",
    )
    add_para(doc, "Three tools. Say them in this order:", bold=True, space_before=4, space_after=4)
    for item in [
        "Next.js is the website — what visitors see.",
        "Sanity is the CMS. Studio is the editor.",
        "Vercel is where it deploys — git push, live URL, production.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "The loop you will show: an editor publishes in Sanity → Next.js pulls the page → Vercel serves it.",
        space_before=4,
    )
    add_para(
        doc,
        "If there is time after the catalog, you name two other live marketing sites: Jamb and Slingshot Bio. "
        "Public pages only. You will not open their CMS.",
        space_before=4,
    )
    add_para(
        doc,
        "Close line (memorize later, with me): The catalog is the product. The CMS is how it changes without a deploy.",
        italic=True,
        space_before=8,
    )

    add_heading(doc, "Not a SaaS")
    add_para(doc, "This is the first idea in the room. If they think you built a login app, the rest of the talk lands wrong.")
    add_table(
        doc,
        ["A normal SaaS", "This site"],
        [
            ["People log in. Dashboard, settings, roles, billing.", "Nobody logs in. The catalog is public."],
            ["The software is the product.", "The catalog markets a real farm’s seed."],
            ["Data belongs to each customer’s account.", "Content is editorial: packets, stories, stock."],
            ["Success is signup, activation, subscription.", "Success is a packet-list request — a lead."],
        ],
    )

    add_heading(doc, "The catalog")
    add_para(
        doc,
        "Thorn & Furrow packs seed March through June. If a variety is listed, they grow it. "
        "Marketing should not wait on engineering for a new heading or a sold-out packet. Those are fields in the CMS.",
    )
    add_para(doc, "What visitors see:", bold=True, space_before=8, space_after=4)
    add_table(
        doc,
        ["Page", "What it is"],
        [
            ["Home  /", "Composed from CMS sections: masthead, featured packet, sowing table, spring list."],
            ["Varieties", "Seventeen packets grouped by family, like a print catalog."],
            ["A packet (e.g. Icehouse Tomato)", "One document: latin name, story, stock, price, photo. Reused on home, index, packet page, and order form."],
            ["Journal", "Seasonal notebook — essays, not a blog calendar."],
            ["Mail order", "A letter, not a cart. Mark packets, tell us your zone, we write back."],
            ["Studio  /studio", "The editor. Desk labeled The farm. Not linked on the public site."],
        ],
    )
    add_para(
        doc,
        "Open if you can:  http://localhost:3000   and   http://localhost:3000/studio",
        italic=True,
        color=MUTED,
    )
    add_para(
        doc,
        "Click like a visitor: homepage, Icehouse Tomato, one journal note, mail order. Then look at Studio — Catalog homepage — the masthead heading. "
        "Do not publish random changes yet. We will do that together.",
        space_before=4,
    )

    add_heading(doc, "How the site is built (plain English)")
    add_para(
        doc,
        "You will not open code in the talk unless they ask. You still need the path, because slide 7 names it.",
    )
    add_para(doc, "1. An editor hits Publish in Sanity Studio.", space_before=6)
    add_para(doc, "2. Next.js asks Sanity for the page using GROQ — Sanity’s query language, not the AI company Groq.")
    add_para(doc, "3. The page is a React Server Component: data is fetched on the server, then HTML goes to the browser. Visitors do not call the CMS from their laptop on every view.")
    add_para(doc, "4. Vercel hosts production. Pages use ISR — they are cached and refresh about every 60 seconds, so the site stays fast. After publish, a signed webhook can drop that cache immediately so we do not wait a full minute.")
    add_para(doc, "5. If Sanity is down, a local backup still renders the catalog. The demo does not go blank. Photographs live in Sanity, with a local file fallback.")
    add_para(doc, "Five kinds of document in this CMS:", bold=True, space_before=10, space_after=4)
    for item in [
        "variety — one packet (Icehouse Tomato)",
        "catalogHomepage — the composed home page",
        "journalEntry — a seasonal essay",
        "workshop — a field day",
        "siteSettings — farm-wide name, nav, footer, SEO",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "Change Icehouse Tomato once. Home, the packet page, and the order form all update, because they point at the same document.",
        space_before=6,
    )

    add_heading(doc, "Words you will hear")
    add_para(doc, "You do not need to write code. You do need to say these in English.")
    add_table(
        doc,
        ["Word", "Say it like this"],
        [
            ["CMS", "The editor marketers use to change the site without asking engineering to ship a new build."],
            ["Studio", "Sanity’s admin screen. Here it lives at /studio. The desk is labeled The farm."],
            ["Document", "One record. Icehouse Tomato is one document — not a Word file."],
            ["Schema / type", "The shape of a document: which fields exist. Engineering owns types. Editors fill fields."],
            ["Deploy", "Putting new code live (git push to Vercel). Changing a heading in Studio is not a deploy."],
            ["GROQ", "How Next.js asks Sanity for content. Not Groq the AI company."],
            ["RSC", "React Server Component — the page runs on the server, then the visitor gets HTML."],
            ["ISR", "Cached pages that refresh on a timer (~60 seconds) instead of rebuilding the whole site."],
            ["Webhook / revalidate", "Sanity pings the site after publish so the new heading can show before the 60-second timer."],
            ["event ID", "A hidden name on a button (hero_order, nav_order). Analytics keys off the ID, not the words on the button."],
            ["GTM", "Google Tag Manager — the tool that listens for those event IDs."],
            ["Lead", "“Please send me a packet list.” Not “create an account.”"],
            ["CDN", "Copies of the site close to visitors, so pages load fast. Vercel does this."],
        ],
    )

    add_heading(doc, "The live demo (four clicks)")
    add_para(doc, "This is the heart of the talk. We will practice it together. Know the shape now:")
    add_table(
        doc,
        ["Step", "You click", "You are showing"],
        [
            ["1", "Homepage /", "A real catalog: masthead, featured packet, sowing table, seventeen packets."],
            ["2", "Icehouse Tomato", "Same document you just saw featured — not a one-off widget."],
            ["3", "Studio → Catalog homepage → change the masthead heading → Publish → reload /", "The heading is live. No deploy. That loop is the work."],
            ["4", "Hover Request a packet list", "The label is copy. The contract is eventId (hero_order / nav_order). Copy can change. Tracking does not break."],
        ],
    )
    add_para(
        doc,
        "If Studio fails on the day: keep going on the public catalog. Say the local seed still renders. Do not freeze.",
        italic=True,
        color=MUTED,
    )

    add_heading(doc, "Two other live sites (if there is time)")
    add_para(
        doc,
        "These are real businesses, shown as public catalogs only. You may name who buys, what they sell, and how a visit becomes a sale. "
        "You may not open Studio, guess field names, or describe the CMS inside. That is an NDA rule — a legal promise not to show private work.",
    )
    add_para(doc, "Jamb  —  https://www.jamb.co.uk/", size=13, bold=True, name="Georgia", space_before=10, space_after=4)
    add_para(
        doc,
        "Pimlico, London (opened 2001). Interior designers, collectors, country-house clients — then US showrooms. "
        "Inventory is two things: unique antiques (chimneypieces, lighting, furniture) and handmade reproductions so a design stays available after the antique sells. "
        "The sale is high-ticket and consultative. The website is the catalog for the showroom, not Amazon checkout. "
        "Sanity owns collections and journal stories. Next.js is the storefront. Click Home, a collection, one product, the journal. Stay on jamb.co.uk.",
    )
    add_para(doc, "Slingshot Bio  —  https://www.slingshotbio.com/", size=13, bold=True, name="Georgia", space_before=10, space_after=4)
    add_para(
        doc,
        "Labs buy this, not shoppers. Flow cytometry: biopharma, CROs, academic cores, cell therapy. "
        "The problem is the control — donor cells expire and drift; plastic beads scatter wrong. "
        "They sell shelf-stable synthetic cell mimics that scatter and stain like real cells. Catalog SKUs plus custom. "
        "A scientist finds a control, reads a protocol, orders a reagent. Repeat purchases. "
        "Sanity owns product and resource documents. Next.js is the catalog and shop. Click Home, How Mimics Work, Shop, Resources. Stay on slingshotbio.com.",
    )
    add_para(
        doc,
        "If they ask to see Jamb or Slingshot Studio: “No. NDA. Thorn & Furrow is the walkthrough.”",
        italic=True,
        space_before=6,
    )

    add_heading(doc, "How the slides fit")
    add_para(doc, "Ten minutes is slides 1–11. Do not skip slides 2 and 3 (not a SaaS, then the stack). If you run long, skip slide 10. Slides 12–17 only if they have time for past work.")
    add_table(
        doc,
        ["Slides", "You are doing"],
        [
            ["1–3", "What this is: marketing site, not SaaS. Next.js, Sanity, Vercel."],
            ["4–8", "The farm, the pages, why this shape, architecture, five documents."],
            ["9", "Live demo — four clicks. This is the talk."],
            ["10–11", "Five decisions. Close line. Stop here if time is up."],
            ["12–17", "Jamb and Slingshot, public only. Then stop."],
        ],
    )

    add_heading(doc, "Rules that do not move")
    for item in [
        "Never put NDA or interview language on the public Thorn & Furrow site. That site should feel like a live catalog.",
        "Never open Jamb or Slingshot Studio, even “just to look.”",
        "Never invent schema or field names for those two sites.",
        "If you do not know, say so and stay on public pages. Do not guess.",
        "You did not have to write this code to present it. You do have to understand this briefing.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "If you want more before we meet")
    add_para(
        doc,
        "Optional. Only if a word still feels empty. You do not need to finish any of these.",
        italic=True,
        color=MUTED,
    )
    add_table(
        doc,
        ["What", "Short read"],
        [
            ["Next.js", "https://nextjs.org/docs — Getting Started / App Router"],
            ["Sanity / Studio", "https://www.sanity.io/docs"],
            ["GROQ", "https://www.sanity.io/docs/content-lake/groq-introduction"],
            ["Vercel", "https://vercel.com/docs"],
            ["React Server Components", "https://react.dev/reference/rsc/server-components"],
            ["ISR", "https://nextjs.org/docs/app/guides/incremental-static-regeneration"],
        ],
    )

    add_heading(doc, "What happens next")
    add_para(
        doc,
        "Bring this file (or just the ideas) when we sit down. We will walk Studio, practice the four clicks, "
        "and then put the speaking script in your hands. Do not memorize a script before we meet — read this instead.",
    )
    add_para(
        doc,
        "The catalog is the product. The CMS is how it changes without a deploy.",
        size=14,
        italic=True,
        name="Georgia",
        space_before=16,
        space_after=4,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")
    if OLD.exists():
        OLD.unlink()
        print(f"Removed {OLD}")


if __name__ == "__main__":
    build()
