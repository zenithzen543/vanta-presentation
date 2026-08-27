"""Mentor pack: a freshman-readable Word guide + a plain-text speaking script."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent
INK = RGBColor(0x12, 0x11, 0x0F)
MOSS = RGBColor(0x3A, 0x4A, 0x38)
MUTED = RGBColor(0x6B, 0x65, 0x60)

BEATS = [
    {
        "clock": "0:00–0:40",
        "slide": "1 · Title",
        "do": None,
        "say": (
            "This is Thorn and Furrow — a live heirloom-seed catalog on Next.js and Sanity. "
            "Product, content model, publish loop. I can open this catalog. "
            "Two other live sites come after the close. I will not open their studios."
        ),
    },
    {
        "clock": "0:40–1:30",
        "slide": "2 · I own the types. They own the catalog.",
        "do": None,
        "say": (
            "The farm packs seed March through June. Catalog Number 14 — if a variety is listed, they grow it. Seventeen packets, no ornamental fillers. "
            "Marketing does not wait on engineering for a new heading or a sold-out packet. Those are fields. "
            "I own the types. They own the weekly catalog. The season is a document, not a deploy."
        ),
    },
    {
        "clock": "1:30–2:10",
        "slide": "3 · What visitors see",
        "do": None,
        "say": (
            "Visitors land on a homepage composed from Sanity sections — masthead, featured packet, sowing table, the spring list. "
            "Varieties is the catalog: seventeen packets grouped by family, the way a print catalog is grouped. "
            "Journal is the seasonal notebook. Mail order is a letter, not a cart: mark packets, tell us your zone, we write back. "
            "Studio exists at /studio for editors. It is not linked on the public site. The storefront should feel live, not like a CMS demo."
        ),
    },
    {
        "clock": "2:10–2:50",
        "slide": "4 · Why this shape",
        "do": None,
        "say": (
            "I did not build a hero, a two-by-three feature grid, and a logo wall. That would look like homework. "
            "A catalog needs four things a marketing site actually needs. "
            "Documents: Icehouse Tomato is one record, reused on home, the index, the packet page, and the order form. "
            "Composition: the homepage is an ordered list of blocks. Editors reorder it without a deploy. "
            "Editor-safe fields: stock, packet price, sowing, isolation, photographs. "
            "Stable tracking: GTM keys off event IDs, not the words on the button."
        ),
    },
    {
        "clock": "2:50–3:35",
        "slide": "5 · Editor to production",
        "do": None,
        "say": (
            "The path is simple. An editor publishes in Sanity Studio. "
            "Next.js pulls the page with GROQ through next-sanity. "
            "Pages are React Server Components in the App Router. "
            "Production uses ISR — pages revalidate about every sixty seconds — plus a signed revalidate API after publish, so we do not wait a full minute. "
            "If Sanity is down, a local seed still renders the catalog. That is resilience, and it is also how a demo does not die. "
            "Photographs come from Sanity assets, with a local file fallback per slug."
        ),
    },
    {
        "clock": "3:35–4:10",
        "slide": "6 · Five documents",
        "do": None,
        "say": (
            "Five document types. Farm settings: name, nav, primary CTA, footer, default SEO. "
            "Catalog homepage: SEO plus an ordered section array. "
            "Variety: latin name, family, days, packet, stock, photo, story. That is the product object. "
            "Journal note and workshop are seasonal content. "
            "The Studio desk is labeled The farm so editors see a catalog, not a generic CMS. "
            "Change Icehouse Tomato once. Home, packet page, and order form all update, because they reference the same document."
        ),
    },
    {
        "clock": "4:10–6:40",
        "slide": "7 · Live demo",
        "do": None,
        "say": (
            "This is the site. Catalog Number 14, Spring 2026. The masthead is a section. "
            "The featured packet is a document — latin name, story, maturity, stock — not a dashboard widget. "
            "The sowing table is editorial and dated for zone 6a. "
            "The spring list is seventeen references to variety documents. Marketing reorders this list. I own the types. "
            "Icehouse Tomato: same document you just saw featured. App Router, static params, ISR. "
            "In Studio, Catalog homepage, Catalog sections, Catalog masthead: I change the Heading and publish. Reload. The heading is live. That publish loop is the work. "
            "Request a packet list: the label is copy. The contract is eventId — hero_order, nav_order. Copy can change. GTM does not."
        ),
    },
    {
        "clock": "6:40–7:20",
        "slide": "8 · What I would defend",
        "do": None,
        "say": (
            "Five decisions I would make again. "
            "Server components and GROQ, not a client-side CMS fetch on every page. "
            "References over duplication. "
            "Event IDs on every CTA so measurement survives a rewrite. "
            "ISR plus a signed revalidate webhook. "
            "No cart. The conversion is a qualified packet-list request — closer to a real marketing funnel than a fake checkout."
        ),
    },
    {
        "clock": "7:20–7:45",
        "slide": "9 · Close",
        "do": None,
        "say": "The catalog is the product. The CMS is how it changes without a deploy.",
    },
    {
        "clock": "7:45–8:00",
        "slide": "10 · Work I cannot open",
        "do": None,
        "say": (
            "Same pattern as an ads business: who buys, what is inventory, how a visit becomes a sale. Then where the CMS sits. "
            "Jamb is a Pimlico dealer. Slingshot Bio sells reagents to flow cytometry labs. "
            "I will not open Studio, schemas, or fields. Public surfaces only."
        ),
    },
    {
        "clock": "8:00–8:35",
        "slide": "11 · Jamb — who buys, what sells",
        "do": None,
        "say": (
            "Jamb opened on Pimlico Road in 2001. The buyer is an interior designer, a collector, a country-house client — London, then Los Angeles, Chicago, Dallas, Atlanta, Palm Beach. "
            "Inventory is two things at once. Unique antiques: chimneypieces from the seventeenth century on, furniture, lighting. "
            "And handmade reproductions from the south London workshop, so a design stays available after the antique sells — that is the original business idea. "
            "The sale is high-ticket and consultative. A designer specs a piece from the catalog. The showroom closes it. Enquiry, not Amazon checkout. "
            "Sanity owns collections and journal stories. Next.js is the storefront."
        ),
    },
    {
        "clock": "8:35–9:05",
        "slide": "12 · Jamb — four public collections",
        "do": None,
        "say": (
            "Four public surfaces, same contract as any marketing catalog. "
            "Fireplaces: antique chimneypieces and stone or marble reproductions, plus a bespoke service. "
            "Lighting: hanging globes, lanterns, wall lights — the Original Globe is the story that started the reproduction line. "
            "Furniture: English country-house seating and tables, antique and made. "
            "Journal: stories for architects and designers, not a blog calendar. "
            "Editors change stock and stories. Engineering owns the types."
        ),
    },
    {
        "clock": "9:05–9:40",
        "slide": "13 · Slingshot — the control problem",
        "do": None,
        "say": (
            "Slingshot Bio sells to flow cytometry labs: biopharma, CROs, academic cores, cell therapy, instrument makers. "
            "The problem is the control. Donor cells expire and drift lot to lot. Polystyrene beads have the wrong scatter. "
            "They sell shelf-stable synthetic cell mimics — polymer particles engineered to scatter, fluoresce, and stain like real cells. "
            "Compensation, unmixing, viability, immunophenotyping, instrument setup. Catalog SKUs plus custom mimics on request. "
            "A scientist finds a control, reads a protocol, orders a reagent. Repeat purchases. "
            "Sanity owns product and resource documents. Next.js is the catalog and shop."
        ),
    },
    {
        "clock": "9:40–10:10",
        "slide": "14 · Slingshot — catalog, argument, shop",
        "do": None,
        "say": (
            "Three public surfaces. Homepage is the catalog of controls. "
            "How Mimics Work is the argument: not a biologic, not a bead. "
            "Shop is the conversion: SKUs grouped by purpose — unmixing, immunophenotyping, CAR-T, custom — then a cart. "
            "Resources sit beside the shop: application notes, data sheets, protocols. That is the education that sells a reagent. "
            "Same rule: I cannot open the desk."
        ),
    },
    {
        "clock": "10:10–10:25",
        "slide": "15 · Stop",
        "do": None,
        "say": (
            "I cannot open those desks. This catalog I can. "
            "Happy to go into schema, caching, or the order form — here."
        ),
    },
]


def set_run(run, *, size=11, bold=False, color=INK, italic=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def add_para(doc, text, *, size=11, bold=False, italic=False, color=INK, name="Calibri", space_before=0, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color, name=name)
    return p


def add_heading(doc, text, level=1):
    return add_para(
        doc,
        text,
        size=18 if level == 1 else 13,
        bold=True,
        color=MOSS if level == 1 else INK,
        name="Georgia",
        space_before=16 if level == 1 else 12,
        space_after=8,
    )


def add_bullet(doc, text, *, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run, size=size)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=INK, size=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        shade_cell(cell, "3A4A38")
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.rows[r].cells[c]
            set_cell_text(cell, value, size=10)
            if r % 2 == 0:
                shade_cell(cell, "F6F4F0")
    doc.add_paragraph()


SLIDES = [
    (
        "1 · Title",
        "Thorn & Furrow is a live heirloom-seed catalog on Next.js and Sanity.",
        "Product, content model, publish loop. This catalog I can open. Two other live sites come later; I will not open their studios.",
        "The deck is Thorn-and-Furrow-Portfolio.pptx — 15 slides. Title names Next.js App Router, Sanity, TypeScript.",
    ),
    (
        "2 · I own the types. They own the catalog.",
        "Engineering owns schemas. Marketing owns the weekly catalog.",
        "Catalog 14, 17 packets grown out, March–June, zone 6a Tivoli. A heading or sold-out packet is a field, not a deploy. The season is a document.",
        "Live URL is on the slide. Studio is not this slide.",
    ),
    (
        "3 · What visitors see",
        "Four public surfaces. Studio is not linked on the storefront.",
        "Home is composed sections. Varieties is the catalog. Journal is field notes. Mail order is a letter, not a cart.",
        "If they ask why no cart: qualified packet-list lead.",
    ),
    (
        "4 · Why this shape",
        "Documents, composition, editor fields, stable tracking.",
        "Icehouse is one record reused everywhere. Homepage is an ordered section array. GTM keys off eventId, not button copy.",
        "If they ask why not a 2×3 grid: that is a SaaS homepage. This is a catalog.",
    ),
    (
        "5 · Editor to production",
        "Studio → GROQ → Next.js → ISR → catalog.",
        "RSC pages. ISR about 60s plus a signed revalidate after publish. Local seed if Sanity is down.",
        "GROQ is Sanity’s query language, not the AI company.",
    ),
    (
        "6 · Five documents",
        "Farm settings, Catalog homepage, Variety, Journal note, Workshop.",
        "Variety is the product object. Change Icehouse once; home, packet page, and order form follow.",
        "Desk is labeled The farm.",
    ),
    (
        "7 · Live demo — this is the talk",
        "Home, Icehouse packet, publish a heading, CTA eventId.",
        "Masthead is a section. Featured packet is a document. Publish is not a deploy. Label is copy; eventId is the contract.",
        "Studio path if they ask: Catalog homepage → Catalog sections → Catalog masthead → Heading → Publish.",
    ),
    (
        "8 · What I would defend",
        "Five decisions: RSC+GROQ, references, eventId, ISR+webhook, no cart.",
        "Server fetch. One document. Tracking survives a rewrite. Fresh after publish. A lead, not a fake checkout.",
        "If Studio failed, local seed still rendered.",
    ),
    (
        "9 · Close",
        "One sentence.",
        "The catalog is the product. The CMS is how it changes without a deploy.",
        "Stop here if they only have ten minutes. Past work is slides 10–15.",
    ),
    (
        "10 · Work I cannot open",
        "Jamb and Slingshot as public businesses. No Studio.",
        "Who buys, what is inventory, how a visit becomes a sale. Then where Sanity sits.",
        "Never invent field names.",
    ),
    (
        "11 · Jamb — who buys, what sells",
        "Designers and collectors. Antiques plus reproductions. Showroom enquiry.",
        "Sanity owns collections and journal stories. Next.js is the storefront. I cannot open the desk.",
        "Screenshot is the public homepage.",
    ),
    (
        "12 · Jamb — four public collections",
        "Fireplaces, lighting, furniture, journal.",
        "Same contract as any marketing catalog. Do not name schemas.",
        "Original Globe is the lighting story that started reproductions.",
    ),
    (
        "13 · Slingshot — the control problem",
        "Labs buy mimics because donor cells expire and beads scatter wrong.",
        "Shelf-stable synthetic cells. Find a control, read a protocol, order a reagent. Repeat purchases.",
        "How Mimics Work is the public argument. Still no Studio.",
    ),
    (
        "14 · Slingshot — catalog, argument, shop",
        "Home, How Mimics Work, shop.",
        "Shop has a cart — that is their conversion, not Thorn & Furrow’s. Products and resources in Sanity.",
        "I cannot open the desk.",
    ),
    (
        "15 · Stop",
        "Questions on this catalog, not those studios.",
        "I cannot open those desks. This catalog I can.",
        "Schema, caching, order form — here.",
    ),
]


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    add_para(doc, "MENTOR PACK  ·  FOR THE PRESENTER", size=10, bold=True, color=MOSS, space_after=4)
    add_para(doc, "Understand the talk, then read the script", size=26, bold=True, name="Georgia", space_after=4)
    add_para(
        doc,
        "This guide is for mentoring. The slides are Thorn-and-Furrow-Portfolio.pptx — 15 slides, that is the deck on the projector. "
        "Give her For-the-Presenter.docx first. Then sit together: Studio, the four-click demo, SCRIPT.txt. "
        "She does not need to have written the code. She does need to understand the story.",
        size=12,
        italic=True,
        color=MUTED,
        name="Georgia",
        space_after=12,
    )
    add_table(
        doc,
        ["File", "What it is"],
        [
            ["Thorn-and-Furrow-Portfolio.pptx", "The deck. 15 slides. Put this on the projector."],
            ["For-the-Presenter.docx", "What she reads before mentoring."],
            ["SCRIPT.txt", "Word-for-word SAY / DO. After she understands the slides."],
            ["Interview-Assistant-Note.txt", "Panel questions and answers. Paste into the interview copilot."],
        ],
    )

    add_heading(doc, "For the mentor (read this first)")
    add_para(
        doc,
        "Coach the opening from the real deck: title, then “I own the types, they own the catalog.” Then what visitors see.",
    )
    for item in [
        "Before Studio: have her say what the catalog is, and what she owns versus marketing, with the Portfolio.pptx slides in front of her.",
        "Then open /studio. Walk Catalog homepage → Catalog sections → Catalog masthead → Heading → Publish → reload /. That loop is the talk.",
        "Have her explain Jamb and Slingshot out loud using only who / inventory / sale. Stop her if she starts inventing field names.",
        "First rehearsal: she reads SCRIPT.txt. Second rehearsal: she uses only Thorn-and-Furrow-Portfolio.pptx. Third: you interrupt with questions.",
        "If she freezes, the rescue line is: “The catalog is the product. The CMS is how it changes without a deploy.”",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Start here — what you are presenting")
    add_para(
        doc,
        "You are not presenting a SaaS product. You are presenting a marketing website.",
        size=14,
        italic=True,
        name="Georgia",
        space_after=10,
    )
    add_para(
        doc,
        "Say this until it is boring: a SaaS is an app people log into (dashboard, settings, billing). "
        "This talk is a public website for a farm that sells seed. Nobody logs in. The pages are the catalog. "
        "Marketers change packets and headings in a CMS. Engineering does not ship a new build for those changes. "
        "At the end you will also name two other live marketing sites (Jamb, Slingshot) that you cannot open inside.",
    )
    add_para(doc, "Your job in the room is three things, in this order:", size=11, bold=True, space_before=8, space_after=4)
    for item in [
        "1. Title and types: live catalog on Next.js and Sanity. I own the types; they own the weekly catalog.",
        "2. Walk Thorn & Furrow live: homepage, Icehouse Tomato, publish a heading, event ID on the button.",
        "3. If there is time: Jamb and Slingshot as public businesses only. No Studio.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "If they ask whether this is a SaaS")
    add_para(doc, "This comparison is not on the deck. Use it only if the panel asks.")
    add_table(
        doc,
        ["A normal SaaS", "This site (marketing catalog)"],
        [
            ["People log in. Dashboard, settings, roles, billing.", "Nobody logs in. The catalog is public."],
            ["The software is the product.", "The catalog is the marketing surface for a real farm."],
            ["Data belongs to each customer’s account.", "Content is editorial: packets, stories, stock."],
            ["Engineering ships features: auth, CRUD, permissions.", "Engineering ships types, pages, publish, SEO, tracking."],
            ["Success is signup, activation, subscription.", "Success is a packet-list request — a marketing lead."],
        ],
    )
    add_para(
        doc,
        "If someone asks “is this like Slack / Notion / a dashboard app?” the answer is no. Those are SaaS. This is the public website that sells the thing.",
        italic=True,
        color=MUTED,
    )

    add_heading(doc, "What the title names")
    add_para(doc, "Slide 1 footer: Next.js App Router, Sanity, TypeScript. Slide 5 is where Studio, GROQ, Next.js, and ISR appear. Vercel is the host of the live catalog URL on slide 2 — not its own slide.")
    add_table(
        doc,
        ["Tool", "What it is", "What you say"],
        [
            ["Next.js", "The website framework (React, App Router).", "Next.js is the site — what visitors see."],
            ["Sanity", "The CMS. Studio is the editor at /studio.", "Sanity is the CMS. Packets are documents, not a Word file."],
            ["TypeScript", "Types shared between Studio fields and pages.", "The types stay honest between CMS and site."],
            ["Vercel", "The host. Git push → live URL, CDN, ISR.", "Vercel is where it deploys. That is production."],
        ],
    )
    add_para(
        doc,
        "The loop, in one line: editor publishes in Sanity → Next.js pulls the page → Vercel serves it. "
        "That is what you are here to show. You do not need to explain React internals.",
        space_before=4,
    )

    add_heading(doc, "The one idea")
    add_para(
        doc,
        "The catalog is the product. The CMS is how it changes without a deploy.",
        size=14,
        italic=True,
        name="Georgia",
        space_after=10,
    )
    add_para(
        doc,
        "In plain English: this is a real-looking website for a farm that sells heirloom seeds. "
        "Marketers can change headings, stock, and stories in Sanity Studio. A programmer does not need to redeploy the site for those changes.",
    )

    add_heading(doc, "What you are showing")
    add_table(
        doc,
        ["Thing", "What it is", "Can you open it?"],
        [
            ["Thorn & Furrow", "A Hudson Valley heirloom-seed catalog. 17 packets, journal, workshops, mail-order letter. Catalog No. 14, Spring 2026.", "Yes. Website and Studio."],
            ["Jamb", "A Pimlico dealer: antique and reproduction fireplaces, lighting, furniture, a journal. jamb.co.uk", "Public website only. No Studio."],
            ["Slingshot Bio", "Lab reagents: synthetic cell mimics for flow cytometry. slingshotbio.com", "Public website only. No Studio."],
        ],
    )

    add_heading(doc, "Hard rules (NDA)")
    add_para(doc, "Jamb and Slingshot were built under NDA. You may talk about the public business. You may not show how they are built inside.")
    for item in [
        "You MAY: name the company, what they sell, who buys, how a visit becomes a sale, and that Sanity feeds public pages (collections/stories, or products/resources).",
        "You MAY NOT: open their Studio, repo, or admin. Do not name schemas, field names, or internal tools.",
        "If they ask “can you show Jamb Studio?” say: “No. NDA. I can name the business and which public surfaces Sanity feeds. Thorn & Furrow is the walkthrough.”",
        "Do not mention NDA, interview, or CMS on the public Thorn & Furrow site. That site should feel like a live catalog.",
        "Do not guess. If you do not know, say you will stay on public surfaces.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Words you need (say them in English first)")
    add_table(
        doc,
        ["Word", "Say it like this"],
        [
            ["CMS", "The editor marketers use to change the website without asking engineering to ship a new build."],
            ["Next.js", "The website framework that draws the public pages. React, App Router."],
            ["Sanity", "The CMS used here. Studio is the editor."],
            ["Vercel", "Where the Next.js site deploys. Git push, live URL, CDN, ISR."],
            ["Studio", "The admin screen of Sanity. For this catalog it lives at /studio. The desk is labeled The farm."],
            ["Document", "One record. Icehouse Tomato is one document. Home, the packet page, and the order form all point at it."],
            ["Section / composition", "The homepage is a list of blocks in order. Editors can reorder them."],
            ["GROQ", "The query language Next.js uses to pull content from Sanity."],
            ["ISR", "Pages refresh about every 60 seconds. A signed revalidate call after publish means we do not wait a full minute."],
            ["eventId", "A hidden name on a button (hero_order, nav_order). Analytics listens to the ID, not the words on the button."],
            ["Mail order / no cart", "Visitors mark packets and tell us their zone. We write back. That is a marketing lead, not a fake checkout."],
            ["Publish vs deploy", "Publish = save live in Studio. Deploy = git push new code. Changing a heading is publish, not deploy."],
            ["Seed fallback", "If Sanity is down, local backup content still renders the catalog. The demo does not go blank."],
        ],
    )

    add_heading(doc, "How the 15 slides fit")
    add_para(doc, "Thorn-and-Furrow-Portfolio.pptx is 15 slides. Ten minutes is 1–9. Slides 10–15 are past work if they have time.")
    add_table(
        doc,
        ["Time", "Slide", "In one sentence"],
        [
            ["0:00–0:40", "1 Title", "Live catalog on Next.js and Sanity. Product, model, publish loop."],
            ["0:40–1:30", "2 I own the types", "They own the weekly catalog. Season is a document."],
            ["1:30–2:10", "3 What visitors see", "Four public pages. No cart. Studio not linked."],
            ["2:10–2:50", "4 Why this shape", "Documents, sections, fields, event IDs."],
            ["2:50–3:35", "5 Editor to production", "Studio → GROQ → Next.js → ISR → catalog."],
            ["3:35–4:10", "6 Five documents", "Change Icehouse once. Every surface follows."],
            ["4:10–6:40", "7 Live demo", "Home, packet, publish, tracking."],
            ["6:40–7:20", "8 Decisions", "Five choices I would make again."],
            ["7:20–7:45", "9 Close", "The memorized line. Stop if time is up."],
            ["7:45–8:00", "10 Work I cannot open", "Jamb and Slingshot. Public only."],
            ["8:00–8:35", "11 Jamb business", "Who, inventory, sale."],
            ["8:35–9:05", "12 Jamb collections", "Fireplaces, lighting, furniture, journal."],
            ["9:05–9:40", "13 Slingshot problem", "The control, then the product."],
            ["9:40–10:10", "14 Slingshot catalog", "Home, argument, shop."],
            ["10:10–10:25", "15 Stop", "Questions on this catalog, not those desks."],
        ],
    )
    add_para(
        doc,
        "If you run long: skip slide 8 (decisions), say the close line, take questions. Past work is 10–15.",
        italic=True,
        color=MUTED,
    )

    add_heading(doc, "Slide by slide — understand, then speak")
    add_para(
        doc,
        "For each slide: first understand what it is arguing. Then read the SAY lines from SCRIPT.txt. "
        "Do not invent extra features.",
        italic=True,
        color=MUTED,
    )
    for num, (title, meaning, gist, coach) in enumerate(SLIDES, start=1):
        add_heading(doc, f"Slide {num}  ·  {title.split(' · ', 1)[-1]}", level=2)
        add_para(doc, "What this slide means", size=10, bold=True, color=MOSS, space_after=2)
        add_para(doc, meaning, space_after=6)
        add_para(doc, "Say it in your own words first", size=10, bold=True, color=MOSS, space_after=2)
        add_para(doc, gist, space_after=6)
        add_para(doc, "Coach note", size=10, bold=True, color=MOSS, space_after=2)
        add_para(doc, coach, italic=True, color=MUTED, space_after=10)

    add_heading(doc, "Live demo — four clicks (practice this)")
    add_para(
        doc,
        "Before you start: catalog at http://localhost:3000 and Studio at http://localhost:3000/studio, already logged in. Hard-refresh once. "
        "Footer should read Content: Sanity live. If it says local seed, Studio edits will not show on the public site.",
    )
    add_table(
        doc,
        ["Step", "You click", "You say"],
        [
            ["1", "Homepage /", "Catalog Number 14. The masthead is a section. The featured packet is a document, not a widget. The sowing table is editorial. The spring list is seventeen references."],
            ["2", "Icehouse Tomato", "Same document you just saw featured."],
            ["3", "Studio → Catalog homepage → Catalog sections → Catalog masthead → Heading → Publish → reload /", "The heading is live. Marketing changed a field. Engineering did not ship a build."],
            ["4", "Hover Request a packet list", "The label is copy. The contract is eventId — hero_order, nav_order. Copy can change. GTM does not."],
        ],
    )
    add_para(doc, "If Studio fails: keep going on the public site. Say the local seed still renders the catalog. Do not apologize for a long time.")

    add_heading(doc, "How to change content in Studio")
    add_para(
        doc,
        "This is the part they will ask her to do live. Coach it until the clicks are muscle memory. "
        "Draft is not live. She must click Publish, then reload the public catalog. That is not a deploy.",
    )
    add_para(doc, "Open http://localhost:3000/studio. Left desk title is The farm.", bold=True, space_before=4)
    add_para(doc, "List: Catalog homepage, Farm settings, then Varieties, Journal, Workshops.")
    add_para(doc, "The demo edit — click by click", size=13, bold=True, name="Georgia", space_before=10, space_after=4)
    add_table(
        doc,
        ["Step", "Click", "Note"],
        [
            ["1", "Catalog homepage (left)", "One document. The homepage is not a Word file."],
            ["2", "Catalog sections", "An ordered list of blocks. Editors can drag to reorder."],
            ["3", "Catalog masthead", "First section on the live home."],
            ["4", "Heading", "Change this. Leave Catalog No and Season unless they ask."],
            ["5", "Publish", "Bottom-right or top. Grey button = a required field is empty."],
            ["6", "Reload http://localhost:3000", "Hard-refresh. The new heading is live. No git push."],
        ],
    )
    add_para(
        doc,
        "What she says while clicking: “Studio is the editor. The desk is labeled The farm so it looks like a catalog, not a generic CMS. I change a field and publish. Next.js pulls it. No deploy.”",
        italic=True,
        color=MUTED,
    )
    add_para(doc, "If they ask her to change something else", size=13, bold=True, name="Georgia", space_before=10, space_after=4)
    add_table(
        doc,
        ["They ask", "Path", "She says"],
        [
            ["Featured packet", "Catalog homepage → Catalog sections → Featured variety → Variety → pick another packet → Publish", "Same homepage document. The feature is a reference, not a paste."],
            ["Reorder the home", "Catalog homepage → Catalog sections → drag blocks → Publish", "Editors compose the page. I own the section types."],
            ["Packet stock or price", "Varieties → Icehouse Tomato → Stock or Packet → Publish", "One document. Home, packet page, and order form all pick it up."],
            ["Packet story / latin", "Same Variety document: Common name, Latin name, Days to maturity, Sowing, Isolation, Kitchen use, Packet photograph", "I do not paste Icehouse in three places."],
            ["Button words", "Catalog masthead → Call to action → Label. Do not touch Analytics event ID (hero_order / nav_order)", "Copy can change. GTM does not."],
            ["Farm name / nav", "Farm settings", "Site chrome. Nav is label plus href."],
            ["Journal", "Journal → open a note → Title / Excerpt / Body → Publish", "Seasonal notebook, not a blog calendar."],
            ["Workshop", "Workshops → open one → Title, Date, Place, Seats, Notes → Publish", "Field days. Same publish loop."],
            ["Add a packet", "Varieties → Create → Common name, Latin name, slug (Generate), family, story → Publish. Then Catalog homepage → Variety index → add reference", "New document, then point the index at it."],
        ],
    )
    add_para(
        doc,
        "Do not type raw HTML. Do not invent new section types. Unknown section types are skipped on the site. "
        "If the public footer still says local seed after Publish, the catalog is not reading Studio — do not claim the edit went live.",
        space_before=4,
    )

    add_heading(doc, "Jamb and Slingshot in four lines each")
    add_para(doc, "Jamb", size=13, bold=True, name="Georgia", space_before=4, space_after=4)
    for item in [
        "Who buys: interior designers, collectors, country-house clients. Pimlico Road plus US showrooms.",
        "What they sell: unique antiques (chimneypieces, lighting, furniture) and handmade reproductions so a design stays available after the antique sells.",
        "How a visit becomes a sale: high-ticket, consultative. The website is the catalog for the showroom, not a cart.",
        "Where Sanity sits: collections and journal stories. Next.js is the storefront.",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Slingshot Bio", size=13, bold=True, name="Georgia", space_before=8, space_after=4)
    for item in [
        "Who buys: flow cytometry labs — biopharma, CROs, academic cores, cell therapy.",
        "What they sell: shelf-stable synthetic cell mimics that scatter and stain like real cells. Catalog SKUs plus custom.",
        "How a visit becomes a sale: a scientist finds a control, reads a protocol, orders a reagent. Repeat purchases.",
        "Where Sanity sits: product and resource documents. Next.js is the catalog and shop.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "If they ask")
    add_table(
        doc,
        ["Question", "Answer in one breath"],
        [
            ["How do you change the homepage?", "Studio → Catalog homepage → Catalog sections → Catalog masthead → Heading → Publish → reload /. No deploy."],
            ["How do you change a packet?", "Varieties → Icehouse Tomato → Stock, Packet, or Story → Publish. One document. Home, packet page, and order form all update."],
            ["Is this a SaaS?", "No. Nobody logs in. It is a public marketing catalog. The conversion is a lead, not a subscription."],
            ["Can you show Jamb or Slingshot Studio?", "No. NDA. I can name the business and which public surfaces Sanity feeds. Thorn & Furrow is the walkthrough."],
            ["How does the homepage work?", "An ordered Sanity array. The UI switches on section type. Editors reorder blocks."],
            ["What is ISR?", "Pages regenerate on a 60-second window, plus a signed revalidate call after publish."],
            ["How do you track CTAs?", "A hidden event ID on the link. GTM listens for the ID, not the label."],
            ["What if Sanity is down?", "Local seed still renders. The catalog does not go blank."],
            ["Is there a cart?", "No. Mail order is a qualified packet-list request — a marketing lead."],
            ["What is a document?", "Icehouse Tomato is one record. Home, packet page, and order form all reference it."],
        ],
    )

    add_heading(doc, "Before you walk in")
    for item in [
        "Memorize the close line.",
        "Do the four demo clicks twice without notes. Then change Icehouse stock once so she feels the same document on three pages.",
        "Read SCRIPT.txt out loud once. Then put it down and use Thorn-and-Furrow-Portfolio.pptx.",
        "If you only have ten minutes: slides 1–9. Skip 8 if long. Past work is 10–15.",
        "Breathe. Point at the screen. Finish the sentence if they interrupt, then answer.",
    ]:
        add_bullet(doc, item)

    add_para(
        doc,
        "The catalog is the product. The CMS is how it changes without a deploy.",
        size=14,
        italic=True,
        name="Georgia",
        space_before=16,
        space_after=8,
    )
    add_para(
        doc,
        "Slides (the file on the projector): Thorn-and-Furrow-Portfolio.pptx. She reads first: For-the-Presenter.docx. Word-for-word lines: SCRIPT.txt. Interview copilot note: Interview-Assistant-Note.txt. Do not use Thorn-and-Furrow-Deck.pptx — that is only a fallback if Portfolio.pptx was locked.",
        size=9,
        color=MUTED,
    )

    path = OUT / "Mentor-Guide.docx"
    doc.save(path)
    return path


def build_script_txt():
    lines = [
        "THORN & FURROW — SPEAKING SCRIPT",
        "Read SAY out loud. Follow DO. Ten minutes is slides 1–9. Slides 10–15 if they have time.",
        "",
        "BEFORE YOU START",
        "- Catalog: http://localhost:3000",
        "- Studio: http://localhost:3000/studio (already logged in)",
        "- Hard-refresh once. If Studio fails, keep going on the public site.",
        "- If you run long, skip slide 8, say the close line, take questions.",
        "",
        "MEMORIZE THIS LINE",
        "The catalog is the product. The CMS is how it changes without a deploy.",
        "",
        "==================================================",
        "",
    ]
    for beat in BEATS:
        lines += [
            f"{beat['clock']}    {beat['slide']}",
            "",
        ]
        if beat.get("do"):
            lines += [f"DO:  {beat['do']}", ""]
        lines += ["SAY:", beat["say"], "", "--------------------------------------------------", ""]
    lines += [
        "END",
        "If they ask about Jamb or Slingshot Studio: No. NDA. Thorn & Furrow is the walkthrough.",
        "",
    ]
    path = OUT / "SCRIPT.txt"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


if __name__ == "__main__":
    print(build_docx())
    print(build_script_txt())
